import streamlit as st
import google.generativeai as genai
from dotenv import load_dotenv
import os
from docx import Document
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors

# Load environment variables
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")
if not api_key:
    st.error("Gemini API Key not found in .env file! Please check your .env file.")
    st.stop()

# Configure Gemini API
try:
    genai.configure(api_key=api_key)
except Exception as e:
    st.error(f"Error configuring Gemini API: {str(e)}")
    st.stop()

# Function to generate a professional resume using Gemini API
def generate_resume(name, jobs, educations, job_type, tone, length, skills):
    try:
        model = genai.GenerativeModel("gemini-2.0-flash")
        # Construct experience string
        experience_str = "\n".join([
            f"- **{job['position']}** at {job['company']}, {job['location']}, {job['date_joined']} to {job['date_left']}\n  - Problems Solved: {job['problems_solved']}\n  - Salary: ${job['salary']:,}\n  - Achievements: [Infer 2-3 achievements based on problems solved]"
            for job in jobs
        ])
        # Construct education string
        education_str = "\n".join([
            f"- **{edu['subject']}**, {edu['institution']} (inferred), {edu['date_joined']} to {edu['completion_date']}, Grade: {edu['grade']}"
            for edu in educations
        ])
        prompt = f"""
        You are an expert resume writer. Create a {length}-length resume for {name} with the following details:
        - Education: {education_str}
        - Experience: {experience_str}
        - Job Type: {job_type}
        - Skills: {skills}
        - Tone: {tone}

        The resume should be highly professional, well-structured, and tailored for the {job_type} role. Include the following sections:
        1. **Personal Information**: Include the name, a professional email (e.g., {name.lower().replace(' ', '.')}@email.com), and a phone number (e.g., +1-555-123-4567).
        2. **Professional Summary**: A concise 3-4 sentence summary highlighting the candidate's experience, skills, and career goals tailored for the {job_type} role.
        3. **Skills**: List the provided skills ({skills}) and infer additional relevant skills for the {job_type} role.
        4. **Experience**: Format each experience entry with the job title, company name, location, date range, problems solved, salary, and 2-3 bullet points detailing achievements inferred from problems solved.
        5. **Education**: Format each education entry with the subject, institution (inferred), date range, and grade.
        6. **Certifications** (optional): Infer relevant certifications for the {job_type} role if applicable.

        Use clear, concise, and action-oriented language. Avoid generic phrases unless supported by specific achievements. Format the resume as plain text with clear section headers (e.g., ### Personal Information) and bullet points for readability.
        """
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Error generating resume: {str(e)}")
        return None

# Function to create a .docx file
def create_docx(resume_text, name):
    doc = Document()
    doc.add_heading(f"{name}'s Resume", 0)
    for line in resume_text.split("\n"):
        if line.startswith("###"):
            doc.add_heading(line.replace("###", "").strip(), level=1)
        elif line.startswith("-"):
            doc.add_paragraph(line.strip(), style="ListBullet")
        else:
            doc.add_paragraph(line.strip())
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# Function to create a .pdf file
def create_pdf(resume_text, name):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 40
    c.setFont("Helvetica-Bold", 16)
    c.drawString(40, y, f"{name}'s Resume")
    y -= 30
    c.setFont("Helvetica", 12)
    for line in resume_text.split("\n"):
        if y < 40:
            c.showPage()
            y = height - 40
        if line.startswith("###"):
            c.setFont("Helvetica-Bold", 14)
            c.drawString(40, y, line.replace("###", "").strip())
            y -= 20
        elif line.startswith("-"):
            c.setFont("Helvetica", 12)
            c.drawString(50, y, f"• {line.replace('-', '').strip()}")
            y -= 15
        else:
            c.setFont("Helvetica", 12)
            c.drawString(40, y, line.strip())
            y -= 15
    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer.getvalue()

# Streamlit UI
def main():
    # Set page configuration
    st.set_page_config(page_title="AI Resume Builder", page_icon="📄", layout="wide")

    # Custom CSS for a modern, beautiful UI
    st.markdown("""
    <style>
    .main {background-color: #f9f9f9;}
    .stButton>button {
        background-color: #1a73e8;
        color: white;
        border-radius: 8px;
        padding: 12px 24px;
        font-weight: bold;
        transition: background-color 0.3s;
    }
    .stButton>button:hover {
        background-color: #1557b0;
    }
    .stTextInput>div>input, .stTextArea>div>textarea, .stSelectbox>div>select {
        border-radius: 8px;
        border: 1px solid #d1d5db;
        padding: 10px;
        background-color: #ffffff;
        box-shadow: 0 1px 2px rgba(0, 0, 0, 0.05);
    }
    .header {
        text-align: center;
        color: #1a73e8;
        font-size: 40px;
        font-weight: bold;
        margin-bottom: 10px;
    }
    .subheader {
        text-align: center;
        color: #4b5563;
        font-size: 18px;
        margin-bottom: 20px;
    }
    .section-header {
        color: #1f2937;
        font-size: 20px;
        font-weight: bold;
        margin-top: 20px;
    }
    .sidebar .stSelectbox {
        margin-bottom: 20px;
    }
    </style>
    """, unsafe_allow_html=True)

    # Header
    st.markdown('<div class="header">AI Resume Builder</div>', unsafe_allow_html=True)
    st.markdown('<div class="subheader">Craft a Professional Resume with AI Precision</div>', unsafe_allow_html=True)
    st.markdown("---")

    # Sidebar for additional settings
    with st.sidebar:
        st.header("Settings")
        download_format = st.selectbox(
            "Download Format",
            ["PDF", "DOCX"],
            help="Choose the format for downloading your resume."
        )
        st.markdown("---")
        st.markdown("Powered by Gemini & Streamlit")

    # Main input form
    st.markdown('<div class="section-header">Tell Us About Yourself</div>', unsafe_allow_html=True)
    with st.form(key="resume_form"):
        # Personal Information
        name = st.text_input(
            "Full Name *",
            placeholder="e.g., John Doe",
            help="Enter your full name as it should appear on the resume."
        )

        # Skills
        skills = st.text_input(
            "Skills (comma-separated) *",
            placeholder="e.g., Python, JavaScript, Project Management",
            help="List your skills relevant to the job role."
        )

        # Job Type, Tone, and Length
        col1, col2, col3 = st.columns(3)
        with col1:
            job_type = st.selectbox(
                "Preferred Job Type *",
                ["Software Engineer", "Data Scientist", "Product Manager", "Graphic Designer", "Other"],
                help="Select the job role you're targeting."
            )
        with col2:
            tone = st.selectbox(
                "Resume Tone *",
                ["Professional", "Friendly", "Creative", "Formal"],
                help="Choose the tone that best fits your industry and personality."
            )
        with col3:
            length = st.selectbox(
                "Resume Length *",
                ["Short (1 page)", "Medium (1-2 pages)", "Detailed (2+ pages)"],
                help="Select the desired length of your resume."
            )

        # Experience Section (Dynamic Jobs)
        st.markdown('<div class="section-header">Work Experience</div>', unsafe_allow_html=True)
        jobs = []
        for i in range(4):  # Allow up to 4 jobs
            with st.expander(f"Job {i+1} (Expand to add details)"):
                company = st.text_input(f"Company Name (Job {i+1})", placeholder="e.g., ABC Corp")
                location = st.text_input(f"Location (Job {i+1})", placeholder="e.g., New York, NY")
                date_joined = st.text_input(f"Date Joined (Job {i+1})", placeholder="e.g., 2022-01")
                date_left = st.text_input(f"Date Left (Job {i+1})", placeholder="e.g., 2023-12 or Present")
                position = st.text_input(f"Position (Job {i+1})", placeholder="e.g., Software Developer")
                problems_solved = st.text_area(f"Problems Solved (Job {i+1})", placeholder="e.g., Optimized database queries...")
                salary = st.text_input(f"Salary (Job {i+1})", placeholder="e.g., 80000")
                if company and location and date_joined and date_left and position and problems_solved and salary:
                    jobs.append({
                        "company": company,
                        "location": location,
                        "date_joined": date_joined,
                        "date_left": date_left,
                        "position": position,
                        "problems_solved": problems_solved,
                        "salary": salary
                    })

        # Education Section (Dynamic Education Entries)
        st.markdown('<div class="section-header">Education</div>', unsafe_allow_html=True)
        educations = []
        for i in range(2):  # Allow up to 2 education entries (e.g., high school and college)
            with st.expander(f"Education {i+1} (Expand to add details)"):
                subject = st.text_input(f"Subject (Education {i+1})", placeholder="e.g., Computer Science")
                institution = st.text_input(f"Institution (Education {i+1})", placeholder="e.g., XYZ University")
                date_joined = st.text_input(f"Date Joined (Education {i+1})", placeholder="e.g., 2018-09")
                completion_date = st.text_input(f"Completion Date (Education {i+1})", placeholder="e.g., 2022-06")
                grade = st.text_input(f"Grade (Education {i+1})", placeholder="e.g., 3.8/4.0")
                if subject and institution and date_joined and completion_date and grade:
                    educations.append({
                        "subject": subject,
                        "institution": institution,
                        "date_joined": date_joined,
                        "completion_date": completion_date,
                        "grade": grade
                    })

        # Submit button
        st.markdown("---")
        submit_button = st.form_submit_button("Generate Resume")

    # Handle form submission
    if submit_button:
        if name and skills and job_type and jobs and educations:
            with st.spinner("Crafting your professional resume..."):
                resume = generate_resume(name, jobs, educations, job_type, tone, length, skills)
                if resume:
                    st.success("Resume generated successfully!")
                    st.subheader("Your Resume")
                    st.markdown(resume)

                    # Prepare download options
                    if download_format == "PDF":
                        resume_file = create_pdf(resume, name)
                        file_extension = "pdf"
                    else:
                        resume_file = create_docx(resume, name)
                        file_extension = "docx"

                    # Download button
                    st.download_button(
                        label=f"Download Resume as {download_format}",
                        data=resume_file,
                        file_name=f"{name}_resume.{file_extension}",
                        mime=f"application/{file_extension}"
                    )
        else:
            st.error("Please fill in all required fields marked with * for at least one job and one education entry.")

    # Footer
    st.markdown("---")
    st.markdown("Generated on March 15, 2025 | Powered by Google Gemini & Streamlit")

if __name__ == "__main__":
    main()